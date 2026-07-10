# -*- coding: utf-8 -*-
"""
生成《项目开发报告》Word文档
校园综合服务平台
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# 标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = '黑体'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading_font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_font.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif i == 2:
        heading_font.size = Pt(14)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(6)
    else:
        heading_font.size = Pt(13)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_paragraph(text, bold=False, size=None, align=None, font_name=None, space_after=None, first_line_indent=None):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_body(text):
    """添加正文段落（首行缩进2字符）"""
    return add_paragraph(text, size=12, first_line_indent=0.74)


def add_bullet(text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p


def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 灰色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()  # 表后空行
    return table


# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_paragraph('校园综合服务平台', bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体')
add_paragraph('项目开发报告', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体')
doc.add_paragraph()
add_paragraph('学    号：____________________', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('姓    名：____________________', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('专    业：____________________', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('指导教师：____________________', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_paragraph('2026年7月', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# 目录页（手动目录）
# ============================================================
add_paragraph('目  录', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体', space_after=12)
toc_items = [
    ('第一章  可行性分析报告', 1),
    ('1.1 项目背景', 2),
    ('1.2 技术可行性', 2),
    ('1.3 经济可行性', 2),
    ('1.4 操作可行性', 2),
    ('1.5 结论', 2),
    ('第二章  需求分析报告', 1),
    ('2.1 用户需求', 2),
    ('2.2 功能需求', 2),
    ('2.3 非功能需求', 2),
    ('第三章  概要设计', 1),
    ('3.1 系统架构设计', 2),
    ('3.2 模块划分', 2),
    ('3.3 接口设计', 2),
    ('第四章  数据库设计', 1),
    ('4.1 数据库概念设计', 2),
    ('4.2 数据库逻辑设计', 2),
    ('4.3 表结构设计', 2),
    ('第五章  详细设计与实现', 1),
    ('5.1 开发环境与工具', 2),
    ('5.2 后端详细设计', 2),
    ('5.3 前端详细设计', 2),
    ('5.4 AI辅助开发说明', 2),
    ('第六章  系统测试', 1),
    ('6.1 测试环境', 2),
    ('6.2 测试用例', 2),
    ('6.3 测试结果与分析', 2),
]
for item, level in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12) if level == 1 else Pt(11)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if level == 1:
        run.bold = True
        p.paragraph_format.space_before = Pt(6)
    else:
        p.paragraph_format.left_indent = Cm(1.5)

doc.add_page_break()

# ============================================================
# 第一章  可行性分析报告
# ============================================================
doc.add_heading('第一章  可行性分析报告', level=1)

doc.add_heading('1.1 项目背景', level=2)
add_body('随着高校信息化建设的不断深入，校园内各类信息服务平台日益增多，但普遍存在功能分散、信息孤岛、使用不便等问题。学生在日常生活中面临以下痛点：')
add_bullet('信息分散：校园公告、失物招领、设施报修等功能分散在不同平台，学生需在多个应用间切换')
add_bullet('失物招领效率低：拾取物品后缺乏统一发布渠道，失主难以快速找到丢失物品')
add_bullet('报修流程繁琐：设施报修需要填写纸质表单或通过电话联系，缺乏流程跟踪和进度反馈')
add_bullet('缺乏一站式服务：缺少整合校园常用服务的综合平台')
add_body('基于上述问题，本项目拟开发一个校园综合服务平台，整合校园公告、失物招领、设施报修、天气查询等核心功能，通过微信小程序和Web管理后台提供一站式校园服务。')

doc.add_heading('1.2 技术可行性', level=2)
add_body('本项目采用前后端分离架构，技术选型如下：')
add_table(
    ['技术', '版本', '用途'],
    [
        ['Spring Boot', '3.2.0', '后端RESTful API框架'],
        ['MyBatis-Plus', '3.5.5', 'ORM数据访问层'],
        ['MySQL', '8.0+', '关系型数据库'],
        ['Redis', '6.x+', '缓存与验证码存储'],
        ['JWT (JJWT)', '0.12.x', 'Token认证'],
        ['Vue 3', '3.x', '管理后台前端框架'],
        ['Element Plus', '最新版', 'UI组件库'],
        ['微信小程序', '原生框架', '学生端移动应用'],
        ['和风天气API', '免费版', '实时天气数据'],
    ]
)
add_body('以上技术均为业界成熟稳定方案，Spring Boot 3.2 提供了完善的微服务支持，MyBatis-Plus 简化了数据访问开发，Vue 3 + Element Plus 提供了高效的前端开发体验。微信小程序原生框架覆盖了绝大多数移动端场景。因此，项目在技术上是可行的。')

doc.add_heading('1.3 经济可行性', level=2)
add_body('本项目为实训教学项目，开发成本主要包括：')
add_bullet('开发工具：VSCode、微信开发者工具等均为免费工具')
add_bullet('运行环境：MySQL Community、Redis均为开源免费软件')
add_bullet('服务器部署：本地开发环境无需额外服务器成本')
add_bullet('第三方服务：和风天气API提供免费开发版本')
add_body('项目无需额外经济投入，开发和运行成本极低，具有很好的经济可行性。')

doc.add_heading('1.4 操作可行性', level=2)
add_body('微信小程序作为学生端入口，无需下载安装，扫码即可使用，符合学生日常使用习惯。管理后台采用Web端，管理员通过浏览器即可访问，操作界面直观友好。系统采用滑块验证码、JWT认证等安全机制，保障系统安全。因此，项目在操作上是可行的。')

doc.add_heading('1.5 结论', level=2)
add_body('综上所述，校园综合服务平台项目在技术、经济、操作三个方面均具备可行性，项目可以实施开发。')

doc.add_page_break()

# ============================================================
# 第二章  需求分析报告
# ============================================================
doc.add_heading('第二章  需求分析报告', level=1)

doc.add_heading('2.1 用户需求', level=2)
add_body('本系统的用户分为两类：学生用户和管理员用户。')
doc.add_heading('2.1.1 学生用户需求', level=3)
add_bullet('浏览校园公告，了解学校最新通知、活动信息')
add_bullet('查看实时天气，方便出行安排')
add_bullet('发布和浏览失物招领信息，提高物品找回率')
add_bullet('在线提交设施报修申请，实时跟踪工单处理进度')
add_bullet('管理个人信息，包括昵称、手机号、学号等')
add_bullet('查看个人发布记录和报修历史')

doc.add_heading('2.1.2 管理员用户需求', level=3)
add_bullet('发布、编辑和管理校园公告')
add_bullet('审核失物招领信息，确保信息真实有效')
add_bullet('受理和分配报修工单，跟踪工单处理状态')
add_bullet('管理系统用户，包括查看和删除用户')
add_bullet('查看系统运营统计数据，如公告数、工单数、用户数')
add_bullet('查看操作日志，进行审计追溯')

doc.add_heading('2.2 功能需求', level=2)
add_body('系统功能分为学生端和管理端两大模块：')

doc.add_heading('2.2.1 学生端功能（微信小程序）', level=3)
add_table(
    ['模块', '功能', '描述'],
    [
        ['首页', '公告轮播', '展示最新校园公告，支持自动轮播切换'],
        ['首页', '天气卡片', '显示实时天气和未来预报'],
        ['首页', '快捷入口', '快速进入失物招领、报修中心等模块'],
        ['首页', '公告列表', '按分类浏览公告，支持分类筛选'],
        ['失物招领', '信息浏览', '列表展示失物招领和寻物启事信息'],
        ['失物招领', '发布信息', '发布失物招领或寻物启事，包含图片上传'],
        ['失物招领', '搜索筛选', '按关键字搜索物品信息'],
        ['失物招领', '详情查看', '查看失物招领详细信息'],
        ['报修中心', '提交报修', '在线填写报修表单，支持图片上传'],
        ['报修中心', '我的报修', '查看个人报修记录和进度'],
        ['个人中心', '用户登录', '微信自动登录，无需手动注册'],
        ['个人中心', '个人信息', '完善和修改个人信息'],
        ['个人中心', '我的发布', '查看个人发布的失物招领'],
    ]
)

doc.add_heading('2.2.2 管理端功能（Web管理后台）', level=3)
add_table(
    ['模块', '功能', '描述'],
    [
        ['仪表盘', '数据概览', '展示系统运营核心数据'],
        ['公告管理', '公告CRUD', '发布、编辑、删除、查看公告'],
        ['失物招领管理', '审核管理', '审核失物招领信息，通过或驳回'],
        ['失物招领管理', '信息管理', '查看和管理已发布信息'],
        ['报修工单管理', '工单处理', '受理、派单、完成、驳回工单'],
        ['报修工单管理', '状态追踪', '实时查看工单处理状态'],
        ['用户管理', '用户管理', '查看用户列表，删除用户'],
        ['数据统计', '统计图表', '查看各模块数据统计'],
        ['系统配置', '配置管理', '管理系统基础配置参数'],
        ['操作日志', '日志查看', '审计系统操作记录'],
    ]
)

doc.add_heading('2.3 非功能需求', level=2)
add_bullet('安全性：密码采用BCrypt加密，认证使用JWT Token，关键操作需管理员权限')
add_bullet('响应速度：接口响应时间不超过2秒，列表查询支持分页')
add_bullet('可用性：系统7×24小时可用，支持50+并发用户')
add_bullet('可维护性：代码规范清晰，数据库设计完整，接口文档完善')
add_bullet('数据一致性：关键操作使用事务管理，逻辑删除而非物理删除')

doc.add_page_break()

# ============================================================
# 第三章  概要设计
# ============================================================
doc.add_heading('第三章  概要设计', level=1)

doc.add_heading('3.1 系统架构设计', level=2)
add_body('系统采用前后端分离架构，整体分为三层：')
add_bullet('前端层：微信小程序（学生端）+ Vue 3管理后台（管理员端）')
add_bullet('后端层：Spring Boot 3.2 RESTful API服务')
add_bullet('数据层：MySQL数据库 + Redis缓存')
add_body('系统架构图如下：')
add_paragraph('┌─────────────────────────────────────────────────────────┐', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('│                    微信小程序（学生端）                    │', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('│   首页   |   失物招领   |   报修中心   |   个人中心       │', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('├─────────────────────────────────────────────────────────┤', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('│                   Vue 3 管理后台（管理员）                │', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('│  仪表盘 | 公告管理 | 失物管理 | 工单管理 | 用户管理      │', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('├─────────────────────────────────────────────────────────┤', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('│              Spring Boot 3.2 RESTful API                 │', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('│  Auth | 公告 | 失物 | 报修 | 文件 | 天气 | 统计 | 配置  │', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('├─────────────────────────────────────────────────────────┤', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('│              MySQL 8.0  |  Redis 6.x                     │', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('└─────────────────────────────────────────────────────────┘', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading('3.2 模块划分', level=2)
add_body('系统按功能划分为以下模块：')

add_paragraph('（1）认证模块', bold=True, size=12)
add_body('负责用户登录、注册、Token签发与校验。学生端通过设备ID自动登录，管理员通过账号密码+滑块验证码登录。采用JWT Token认证，支持@NoAuth和@RoleRequired注解实现细粒度权限控制。')

add_paragraph('（2）公告模块', bold=True, size=12)
add_body('提供校园公告的发布、查询、编辑、删除功能。支持分类筛选和全文搜索，公告列表按创建时间降序排列。')

add_paragraph('（3）失物招领模块', bold=True, size=12)
add_body('支持失物招领和寻物启事两类信息的发布与浏览，管理员审核机制保障信息真实性。支持按类型、状态、关键字搜索。')

add_paragraph('（4）报修工单模块', bold=True, size=12)
add_body('支持在线提交报修申请，工单状态机流转（待处理→处理中→已完成/已驳回），支持图片上传、处理人分配、进度追踪。')

add_paragraph('（5）文件上传模块', bold=True, size=12)
add_body('支持单图和多图上传，限制文件类型和大小，返回图片URL供其他模块使用。')

add_paragraph('（6）天气查询模块', bold=True, size=12)
add_body('对接和风天气API，提供实时天气查询和天气预报功能，支持Redis缓存减少API调用。')

add_paragraph('（7）统计模块', bold=True, size=12)
add_body('聚合展示系统运营数据，包括公告数、工单数、用户数等核心指标。')

add_paragraph('（8）系统配置模块', bold=True, size=12)
add_body('管理系统配置参数，支持配置项的增删改查。')

doc.add_heading('3.3 接口设计', level=2)
add_body('后端采用RESTful API设计风格，统一使用 /api/ 前缀，返回统一的Result<T>响应格式。')

add_table(
    ['接口分类', '路径', '说明'],
    [
        ['认证', '/api/auth/**', '登录、注册、Token管理'],
        ['验证码', '/api/captcha/**', '滑块验证码生成与验证'],
        ['公告', '/api/announcement/**', '公告CRUD'],
        ['失物招领', '/api/lost-found/**', '失物信息CRUD与审核'],
        ['报修工单', '/api/repair/**', '工单CRUD与状态流转'],
        ['文件上传', '/api/file/**', '单图/多图上传'],
        ['天气', '/api/weather/**', '实时天气与预报查询'],
        ['统计', '/api/statistics', '系统运营数据统计'],
        ['用户管理', '/api/users/**', '用户CRUD'],
        ['系统配置', '/api/system-config/**', '配置项管理'],
        ['操作日志', '/api/operation-logs/**', '操作日志查询'],
    ]
)

add_body('统一响应格式：')
add_paragraph('{ "code": 200, "msg": "操作成功", "data": {...} }', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_body('业务状态码说明：200-成功，400-参数错误，401-未登录，403-权限不足，404-资源不存在，500-服务器错误。')

doc.add_page_break()

# ============================================================
# 第四章  数据库设计
# ============================================================
doc.add_heading('第四章  数据库设计', level=1)

doc.add_heading('4.1 数据库概念设计', level=2)
add_body('系统包含7张核心数据表，实体关系如下：')
add_bullet('用户表（user）：存储学生和管理员信息，是系统的核心实体')
add_bullet('公告表（announcement）：存储校园公告，关联用户表的发布人')
add_bullet('失物信息表（lost_found）：存储失物招领和寻物启事，关联用户表的发布人和审核人')
add_bullet('报修工单表（repair_order）：存储设施报修工单，关联用户表的提交人和处理人')
add_bullet('工单处理记录表（repair_log）：记录工单状态变更历史，关联工单表')
add_bullet('操作日志表（operation_log）：记录系统操作日志，关联用户表')
add_bullet('系统配置表（system_config）：存储系统配置参数，独立实体')

add_body('数据库名称为 campus，字符集采用 utf8mb4，存储引擎为 InnoDB，支持事务和外键逻辑关联。')

doc.add_heading('4.2 数据库逻辑设计', level=2)
add_body('数据库共设计7张表、22个索引（含主键、唯一索引、普通索引、复合索引、全文索引），覆盖所有高频查询场景。关键索引设计策略如下：')
add_bullet('复合索引优先：idx_status_created（工单列表按状态+时间排序）、idx_type_status（失物按类型+状态筛选）、idx_module_action（日志按模块+操作查询）')
add_bullet('全文索引：ft_title_content（公告标题+内容全文搜索）')
add_bullet('唯一索引：uk_open_id（微信登录唯一标识）、uk_order_no（工单编号唯一）')
add_bullet('逾期提醒索引：idx_estimated_complete（报修工单预计完成时间）')

doc.add_heading('4.3 表结构设计', level=2)

add_paragraph('4.3.1 用户表（user）', bold=True, size=12)
add_table(
    ['字段名', '类型', '说明'],
    [
        ['id', 'BIGINT', '主键ID，自增'],
        ['open_id', 'VARCHAR(64)', '微信OpenID，唯一标识'],
        ['email', 'VARCHAR(128)', '邮箱'],
        ['password', 'VARCHAR(255)', '密码(BCrypt加密)'],
        ['nickname', 'VARCHAR(64)', '昵称'],
        ['avatar', 'VARCHAR(255)', '头像URL'],
        ['student_no', 'VARCHAR(32)', '学号'],
        ['phone', 'VARCHAR(20)', '手机号'],
        ['role', 'TINYINT', '角色：0-学生, 1-管理员'],
        ['status', 'TINYINT', '状态：0-正常, 1-禁用'],
        ['created_at', 'DATETIME', '创建时间'],
        ['updated_at', 'DATETIME', '更新时间'],
        ['is_deleted', 'TINYINT', '逻辑删除标志'],
    ]
)

add_paragraph('4.3.2 公告表（announcement）', bold=True, size=12)
add_table(
    ['字段名', '类型', '说明'],
    [
        ['id', 'BIGINT', '主键ID，自增'],
        ['title', 'VARCHAR(128)', '公告标题'],
        ['content', 'TEXT', '公告内容'],
        ['category', 'VARCHAR(32)', '分类：教务通知/活动通知/紧急通知'],
        ['publisher_id', 'BIGINT', '发布人ID'],
        ['view_count', 'INT', '浏览量'],
        ['status', 'TINYINT', '状态：0-草稿, 1-已发布'],
        ['created_at', 'DATETIME', '创建时间'],
        ['updated_at', 'DATETIME', '更新时间'],
        ['is_deleted', 'TINYINT', '逻辑删除标志'],
    ]
)

add_paragraph('4.3.3 失物信息表（lost_found）', bold=True, size=12)
add_table(
    ['字段名', '类型', '说明'],
    [
        ['id', 'BIGINT', '主键ID，自增'],
        ['type', 'TINYINT', '类型：0-失物招领, 1-寻物启事'],
        ['item_name', 'VARCHAR(128)', '物品名称'],
        ['category', 'VARCHAR(32)', '物品类别'],
        ['images', 'VARCHAR(500)', '图片URL列表(JSON数组)'],
        ['description', 'TEXT', '详细描述'],
        ['location', 'VARCHAR(128)', '拾取/丢失地点'],
        ['contact_person', 'VARCHAR(32)', '联系人姓名'],
        ['contact_phone', 'VARCHAR(20)', '联系电话'],
        ['status', 'TINYINT', '状态：0-待审核, 1-已发布, 2-未通过, 3-已结束'],
        ['user_id', 'BIGINT', '发布人ID'],
        ['auditor_id', 'BIGINT', '审核人ID'],
        ['audit_time', 'DATETIME', '审核时间'],
        ['reject_reason', 'VARCHAR(255)', '驳回原因'],
        ['created_at', 'DATETIME', '创建时间'],
        ['updated_at', 'DATETIME', '更新时间'],
        ['is_deleted', 'TINYINT', '逻辑删除标志'],
    ]
)

add_paragraph('4.3.4 报修工单表（repair_order）', bold=True, size=12)
add_table(
    ['字段名', '类型', '说明'],
    [
        ['id', 'BIGINT', '主键ID，自增'],
        ['order_no', 'VARCHAR(32)', '工单编号'],
        ['repair_type', 'VARCHAR(32)', '报修类型'],
        ['campus', 'VARCHAR(32)', '校区'],
        ['building', 'VARCHAR(32)', '楼栋'],
        ['room', 'VARCHAR(32)', '房间号'],
        ['description', 'TEXT', '故障描述'],
        ['images', 'VARCHAR(500)', '图片URL列表(JSON数组)'],
        ['contact_person', 'VARCHAR(32)', '联系人姓名'],
        ['contact_phone', 'VARCHAR(20)', '联系电话'],
        ['status', 'TINYINT', '状态：0-待处理, 1-处理中, 2-已完成, 3-已驳回'],
        ['reject_reason', 'VARCHAR(255)', '驳回原因'],
        ['handle_result', 'VARCHAR(255)', '处理结果'],
        ['user_id', 'BIGINT', '提交人ID'],
        ['handler_id', 'BIGINT', '处理人ID'],
        ['estimated_complete_time', 'DATETIME', '预计完成时间'],
        ['created_at', 'DATETIME', '提交时间'],
        ['handle_time', 'DATETIME', '受理时间'],
        ['complete_time', 'DATETIME', '完成时间'],
        ['updated_at', 'DATETIME', '更新时间'],
        ['is_deleted', 'TINYINT', '逻辑删除标志'],
    ]
)

add_paragraph('4.3.5 其他表', bold=True, size=12)
add_table(
    ['表名', '核心字段', '说明'],
    [
        ['repair_log', 'order_id, action, from_status, to_status', '工单状态变更记录'],
        ['operation_log', 'user_id, module, action, ip_address', '系统操作审计日志'],
        ['system_config', 'config_key, config_value', '系统配置参数'],
    ]
)

doc.add_page_break()

# ============================================================
# 第五章  详细设计与实现
# ============================================================
doc.add_heading('第五章  详细设计与实现', level=1)

doc.add_heading('5.1 开发环境与工具', level=2)
add_table(
    ['类别', '工具/技术', '版本'],
    [
        ['操作系统', 'Windows 11', '-'],
        ['JDK', 'Oracle JDK', '17+'],
        ['后端框架', 'Spring Boot', '3.2.0'],
        ['ORM框架', 'MyBatis-Plus', '3.5.5'],
        ['数据库', 'MySQL', '8.0+'],
        ['缓存', 'Redis', '6.x+'],
        ['前端框架', 'Vue 3 + Element Plus', '-'],
        ['小程序框架', '微信原生框架', '-'],
        ['开发工具', 'VSCode + 微信开发者工具', '-'],
        ['AI辅助工具', 'Trae IDE (DeepSeek-V4-Flash)', '-'],
        ['版本管理', 'Git + GitHub', '-'],
        ['构建工具', 'Maven + npm', '-'],
    ]
)

doc.add_heading('5.2 后端详细设计', level=2)

add_paragraph('5.2.1 项目结构', bold=True, size=12)
add_paragraph('campus-backend/', size=10)
add_paragraph('├── src/main/java/com/campus/', size=10)
add_paragraph('│   ├── annotation/     # 自定义注解（@NoAuth, @RoleRequired）', size=10)
add_paragraph('│   ├── common/         # 通用组件（Result, JwtUtil, 拦截器, 异常处理, 自动填充）', size=10)
add_paragraph('│   ├── config/         # 配置类（CORS, JWT, MyBatis-Plus, Redis）', size=10)
add_paragraph('│   ├── controller/     # API控制器（11个）', size=10)
add_paragraph('│   ├── dto/           # 数据传输对象', size=10)
add_paragraph('│   ├── entity/        # 数据实体（7个）', size=10)
add_paragraph('│   ├── enums/         # 枚举类（5个）', size=10)
add_paragraph('│   ├── mapper/        # MyBatis-Plus数据访问层', size=10)
add_paragraph('│   └── service/       # 业务逻辑层（接口+实现）', size=10)
add_paragraph('└── src/main/resources/', size=10)
add_paragraph('    ├── application.yml # 主配置', size=10)
add_paragraph('    └── mapper/        # XML映射文件', size=10)

add_paragraph('5.2.2 认证模块实现', bold=True, size=12)
add_body('认证模块是整个系统的安全基础，采用JWT（JSON Web Token）实现无状态认证。')
add_body('JwtConfig 类通过 @ConfigurationProperties(prefix = "jwt") 绑定配置，采用构造函数绑定模式，在校验密钥长度（≥32字节）和过期时间（>0）时启动 Fail-Fast，配置错误则项目无法启动，从源头避免安全隐患。密钥通过环境变量 JWT_SECRET 注入，生产环境与开发环境使用不同的密钥。')
add_body('JwtUtil 基于 JJWT 0.12.x 库实现 Token 的生成与解析，Token 中包含 userId（subject）和 role（自定义 claim），使用 HMAC-SHA256 算法签名。')
add_body('JwtAuthInterceptor 实现 HandlerInterceptor 接口，在 preHandle 方法中完成认证鉴权：首先匹配白名单路径放行，然后检查方法/类上的 @NoAuth 和 @RoleRequired 注解，解析 Authorization Header 中的 Bearer Token，将 userId/role 设置到 UserContext 中。在 afterCompletion 方法中清除 UserContext，防止 ThreadLocal 内存泄漏。')
add_body('权限模型采用层级设计：role 值越大权限越高，@RoleRequired(N) 允许 role ≥ N 的用户访问。例如 @RoleRequired(1) 表示只有管理员（role=1）可访问，@RoleRequired(0) 表示所有认证用户均可访问。')

add_paragraph('5.2.3 报修工单模块实现', bold=True, size=12)
add_body('报修工单模块是系统核心业务模块，采用状态机模式管理工单流转。')
add_body('RepairOrderServiceImpl 中定义了状态流转映射表 VALID_TRANSITIONS，明确定义合法流转路径：待处理（PENDING）→ 处理中（PROCESSING）/ 已驳回（REJECTED），处理中（PROCESSING）→ 已完成（COMPLETED）/ 已驳回（REJECTED），已驳回（REJECTED）→ 待处理（PENDING），已完成（COMPLETED）为终态。')
add_body('工单编号采用 "R" + yyyyMMdd + 6位随机数 的格式生成，如 R20260701001。为保证唯一性，在 createOrder 方法中实现了重试机制：最多尝试3次生成随机编号，捕获 DuplicateKeyException 后自动重试。')
add_body('关键业务方法使用 @Transactional 注解保证原子性，包括 createOrder（创建工单+记录日志）和 updateStatus（更新状态+记录日志）。')

add_paragraph('5.2.4 失物招领模块实现', bold=True, size=12)
add_body('失物招领模块支持两种类型（失物招领/寻物启事），采用审核机制确保信息真实性。')
add_body('LostFoundController 实现了角色权限区分：管理员可查看全部记录，普通用户仅能查看已发布和已结束的记录。管理员审核接口（PUT /api/lost-found/{id}/audit）自动设置 auditorId 和 auditTime，支持填写驳回原因。')
add_body('updateById 方法使用 LambdaUpdateWrapper 限定可更新字段白名单，防止全字段更新带来的安全风险。列表查询支持按类型、状态、关键字（物品名/地点/描述）和用户ID进行多条件组合筛选。')

add_paragraph('5.2.5 滑块验证码模块实现', bold=True, size=12)
add_body('CaptchaService 使用纯 Java AWT 图形绘制技术实现滑块验证码，无需依赖第三方验证码服务。')
add_body('在 280×150 像素的背景图上绘制渐变背景和随机干扰线，在随机位置抠出 40×40 像素的拼图块（添加白边和半透明效果，原图挖空区域变暗），返回 base64 编码的背景图和拼图块。')
add_body('验证码的正确位置存入 Redis（5分钟有效），验证时容差 5 像素。验证通过后生成一次性 passToken（5分钟有效），用于后续的登录或密码重置操作，防止重放攻击。')

add_paragraph('5.2.6 统一异常处理', bold=True, size=12)
add_body('GlobalExceptionHandler 使用 @RestControllerAdvice 实现全局异常统一处理，覆盖以下异常类型：参数校验异常（MethodArgumentNotValidException）、业务异常（BusinessException）、JWT过期异常、JWT签名异常、参数非法异常、空指针异常、数字格式异常等。对于未预料的异常，返回通用错误信息，避免内部信息泄露。')

add_paragraph('5.2.7 缓存设计', bold=True, size=12)
add_body('RedisConfig 配置了 Jackson2JsonRedisSerializer 作为序列化方案，禁用了 Jackson 的 activateDefaultTyping 以防止反序列化 RCE 漏洞。配置了 RedisCacheManager 支持 Spring @Cacheable/@CacheEvict 注解缓存。')
add_body('UserServiceImpl 中，getById 方法使用 @Cacheable 缓存用户查询结果，缓存前将密码字段置空防止泄露；updateById 和 removeById 使用 @CacheEvict 在数据变更后清除缓存，保证数据一致性。')

doc.add_heading('5.3 前端详细设计', level=2)

add_paragraph('5.3.1 微信小程序前端', bold=True, size=12)
add_body('小程序采用原生框架开发，包含4个主页面：首页、失物招领、报修中心、个人中心。')
add_body('首页集成了公告轮播（支持自动轮播）、实时天气卡片（调用和风天气API）、快捷功能入口和公告分类列表。失物招领页面采用服务端分页，支持触底加载更多，减少一次性数据传输。报修中心支持在线提交报修表单，包含报修类型选择、校区/楼栋/房间号选择、故障描述和图片上传。')
add_body('请求封装在 utils/request.js 中，统一管理 baseUrl、请求头、Token 注入和错误处理。使用 loadingCount 计数器模式确保 wx.showLoading/wx.hideLoading 配对使用，避免并发请求时的加载状态混乱。')
add_body('登录采用设备ID持久化方案：在 app.js 的 onLaunch 中生成唯一 deviceId 并存入本地存储，每次启动时使用该 deviceId 调用 wx-login 接口实现自动登录，同一设备始终使用同一用户标识。')

add_paragraph('5.3.2 Vue 3 管理后台前端', bold=True, size=12)
add_body('管理后台基于 Vue 3 + Element Plus 开发，采用 Vite 构建工具，使用 Vue Router 实现路由管理。')
add_body('页面包括：登录页（含滑块验证码组件）、仪表盘（数据概览）、公告管理（列表+表单）、失物招领管理（审核+管理）、报修工单管理（受理+状态更新）、用户管理、数据统计、系统配置、操作日志等。')
add_body('API 请求封装在 api/index.js 中，使用 axios 创建实例，统一配置 baseURL、超时时间和响应拦截器。响应拦截器统一处理业务逻辑错误提示、401跳转登录、403权限不足等场景。')
add_body('CaptchaSlider 组件实现了滑块验证码的完整交互流程：获取验证码→渲染背景图和拼图块→拖拽拼图→验证位置→返回 passToken。')

doc.add_heading('5.4 AI辅助开发说明', level=2)
add_body('本项目在开发过程中广泛使用了 Trae IDE 内置的 AI 编程助手（基于 DeepSeek-V4-Flash 模型），以下是在各模块开发中使用 AI 辅助的具体情况。')

add_paragraph('5.4.1 AI辅助开发方式', bold=True, size=12)
add_body('在项目开发过程中，主要采用以下方式使用 AI 辅助编程：')
add_bullet('代码生成：通过自然语言描述功能需求，由 AI 生成对应的代码实现')
add_bullet('代码审查：将已编写的代码提交给 AI 审查，发现潜在问题')
add_bullet('Bug修复：向 AI 描述错误现象和日志，由 AI 定位问题并给出修复方案')
add_bullet('架构咨询：在技术选型和方案设计阶段向 AI 咨询最佳实践')
add_bullet('文档编写：使用 AI 辅助编写项目文档和注释')

add_paragraph('5.4.2 后端模块AI辅助实现', bold=True, size=12)
add_body('（1）认证模块：AI 辅助生成了 JWT 认证拦截器的完整实现，包括 Token 解析、白名单管理、@NoAuth/@RoleRequired 注解解析逻辑。特别是在处理 @NoAuth 和 @RoleRequired 注解冲突检测时，AI 提供了使用 AnnotationUtils.findAnnotation 进行方法级/类级注解查找的实现方案，以及拦截器中 UserContext 的 ThreadLocal 管理机制。')
add_body('（2）报修工单模块：AI 辅助设计了工单状态机模式，定义了 VALID_TRANSITIONS 合法流转路径映射表，并生成了工单编号自动生成逻辑（含重试机制）。在 @Transactional 事务管理方面，AI 建议将 createOrder 和 updateStatus 方法独立封装事务，确保工单创建和日志记录的一致性。')
add_body('（3）滑块验证码模块：AI 辅助实现了基于 Java AWT 的纯图形验证码生成，包括渐变背景绘制、随机干扰线生成、拼图块抠取和半透明效果实现。Redis 存储验证码位置和 passToken 的一次性使用策略也是 AI 推荐的方案。')
add_body('（4）缓存设计：AI 辅助配置了 Jackson2JsonRedisSerializer 序列化方案，并特别提醒了 activateDefaultTyping 的 RCE 安全风险，建议使用 Jackson2JsonRedisSerializer 替代 GenericJackson2JsonRedisSerializer。')

add_body('以下为AI辅助开发过程中的典型提示词示例：')
add_paragraph('提示词示例1：生成JWT认证拦截器', bold=True, size=11)
add_paragraph('"请实现一个Spring Boot的HandlerInterceptor，用于JWT Token认证拦截。需要支持：1）白名单路径放行；2）@NoAuth注解标记免登录接口；3）@RoleRequired注解实现角色权限控制；4）从Authorization Header解析Bearer Token；5）将userId和role设置到请求上下文UserContext中；6）请求结束后清除ThreadLocal。"', size=10, first_line_indent=0.74)
add_paragraph('AI响应：生成了JwtAuthInterceptor完整代码，包含preHandle和afterCompletion方法，实现了白名单匹配、注解解析、Token校验、UserContext设置等完整逻辑。', size=10, first_line_indent=0.74)

add_paragraph('提示词示例2：设计工单状态机', bold=True, size=11)
add_paragraph('"请设计一个报修工单的状态机实现。工单状态有：PENDING(待处理)、PROCESSING(处理中)、COMPLETED(已完成)、REJECTED(已驳回)。合法流转：PENDING→PROCESSING/REJECTED, PROCESSING→COMPLETED/REJECTED, REJECTED→PENDING, COMPLETED为终态。需要事务支持，并记录操作日志。"', size=10, first_line_indent=0.74)
add_paragraph('AI响应：生成了RepairOrderServiceImpl的updateStatus方法，包含状态映射表VALID_TRANSITIONS、状态流转合法性校验、@Transactional事务管理，以及根据目标状态自动设置处理时间/处理人/完成时间等字段的逻辑。', size=10, first_line_indent=0.74)

add_paragraph('5.4.3 前端模块AI辅助实现', bold=True, size=12)
add_body('（1）微信小程序：AI 辅助生成了首页公告轮播组件、天气卡片展示、失物招领列表（服务端分页+触底加载）、报修表单（含图片上传）等核心页面代码。特别在请求封装 utils/request.js 中，AI 提供了 loadingCount 计数器模式防止 wx.showLoading/wx.hideLoading 未配对的问题。')
add_body('（2）管理后台：AI 辅助生成了 Vue 3 组件的完整代码，包括 CaptchaSlider 滑块验证码组件、Login 登录页（含验证码验证流程）、Dashboard 仪表盘（数据统计展示）、各管理页面（CRUD 表格+表单）等。')
add_body('（3）API封装：AI 辅助设计了 axios 实例的响应拦截器，统一处理 Token 过期自动跳转登录、业务错误提示、网络异常处理等场景。')

add_paragraph('提示词示例3：生成小程序请求封装', bold=True, size=11)
add_paragraph('"请封装微信小程序的网络请求工具函数。需要支持：1）统一baseUrl配置；2）自动注入Authorization请求头；3）统一错误处理，401时自动跳转登录；4）提供get/post/put/del便捷方法；5）loading状态管理，防止并发请求时showLoading/hideLoading未配对。"', size=10, first_line_indent=0.74)
add_paragraph('AI响应：生成了utils/request.js文件，包含全局loadingCount计数器、request核心方法、get/post/put/del快捷导出、401自动重登录逻辑。', size=10, first_line_indent=0.74)

add_paragraph('5.4.4 AI辅助Bug修复案例', bold=True, size=12)
add_body('在项目开发过程中，AI 辅助定位并修复了多个关键 Bug：')
add_bullet('管理后台白屏问题：AI 分析发现 api/index.js 缺少命名导出（authApi、captchaApi 等），导致 Vue Router 加载模块失败，添加导出后修复')
add_bullet('小程序重复创建用户：AI 定位到 app.js 使用 wx.login 临时 code 作为登录凭证，每次 code 不同导致每次创建新用户，改为 deviceId 持久化方案修复')
add_bullet('Redis MISCONF 错误：AI 分析日志发现 Redis 无法持久化 RDB 快照，给出 config set stop-writes-on-bgsave-error no 修复命令')
add_bullet('枚举序列化问题：AI 发现工单状态在管理后台不显示的原因为枚举缺少 @JsonValue/@JsonCreator 注解，添加后修复')

add_paragraph('5.4.5 AI辅助开发总结', bold=True, size=12)
add_body('通过使用 AI 辅助编程，项目开发效率得到显著提升。在代码生成方面，AI 能够快速生成符合规范的 Spring Boot 后端代码和 Vue 3/小程序前端代码。在问题排查方面，AI 能够通过分析日志和代码快速定位问题根因，提供准确的修复方案。在架构设计方面，AI 提供了技术选型建议和最佳实践参考。')
add_body('但同时也需要注意的是，AI 生成的代码需要人工审查和测试验证，特别是在安全性和事务一致性方面，需要开发者具备足够的判断能力。AI 辅助编程是提效工具，不能完全替代开发者的技术决策。')

doc.add_page_break()

# ============================================================
# 第六章  系统测试
# ============================================================
doc.add_heading('第六章  系统测试', level=1)

doc.add_heading('6.1 测试环境', level=2)
add_table(
    ['环境', '配置'],
    [
        ['服务器', '本地开发环境（Windows 11）'],
        ['数据库', 'MySQL 8.0 + Redis 6.x'],
        ['后端', 'Spring Boot 3.2 (JDK 17)'],
        ['前端', 'Vue 3 + 微信小程序'],
        ['浏览器', 'Chrome 最新版'],
        ['网络', 'localhost 本地网络'],
    ]
)

doc.add_heading('6.2 测试用例', level=2)

add_paragraph('6.2.1 认证模块测试', bold=True, size=12)
add_table(
    ['编号', '测试用例', '预期结果', '实际结果'],
    [
        ['TC-AUTH-01', '管理员正确登录', '返回Token和用户信息', '通过'],
        ['TC-AUTH-02', '管理员错误密码登录', '返回401错误提示', '通过'],
        ['TC-AUTH-03', '未登录访问需认证接口', '返回401未登录', '通过'],
        ['TC-AUTH-04', '使用过期Token访问', '返回401 Token过期', '通过'],
        ['TC-AUTH-05', '滑块验证码验证通过', '返回passToken', '通过'],
        ['TC-AUTH-06', '滑块验证码位置错误', '返回验证失败', '通过'],
        ['TC-AUTH-07', '学生无权限访问管理接口', '返回403权限不足', '通过'],
        ['TC-AUTH-08', '退出登录', '清除登录状态', '通过'],
    ]
)

add_paragraph('6.2.2 公告模块测试', bold=True, size=12)
add_table(
    ['编号', '测试用例', '预期结果', '实际结果'],
    [
        ['TC-ANN-01', '游客查看公告列表', '返回已发布公告列表', '通过'],
        ['TC-ANN-02', '按分类筛选公告', '返回对应分类公告', '通过'],
        ['TC-ANN-03', '查看公告详情', '返回公告完整内容', '通过'],
        ['TC-ANN-04', '管理员创建公告', '公告创建成功', '通过'],
        ['TC-ANN-05', '管理员编辑公告', '公告更新成功', '通过'],
        ['TC-ANN-06', '管理员删除公告', '公告逻辑删除', '通过'],
        ['TC-ANN-07', '学生创建公告', '返回403权限不足', '通过'],
    ]
)

add_paragraph('6.2.3 失物招领模块测试', bold=True, size=12)
add_table(
    ['编号', '测试用例', '预期结果', '实际结果'],
    [
        ['TC-LF-01', '游客查看失物列表', '返回已发布/已结束记录', '通过'],
        ['TC-LF-02', '按类型筛选失物信息', '返回对应类型信息', '通过'],
        ['TC-LF-03', '按关键字搜索失物', '返回匹配结果', '通过'],
        ['TC-LF-04', '查看失物详情', '返回完整信息', '通过'],
        ['TC-LF-05', '学生发布失物信息', '创建成功，状态为待审核', '通过'],
        ['TC-LF-06', '管理员审核通过', '状态变为已发布', '通过'],
        ['TC-LF-07', '管理员审核驳回', '状态变为未通过，含驳回原因', '通过'],
        ['TC-LF-08', '非管理员查看待审核记录', '不返回待审核记录', '通过'],
    ]
)

add_paragraph('6.2.4 报修工单模块测试', bold=True, size=12)
add_table(
    ['编号', '测试用例', '预期结果', '实际结果'],
    [
        ['TC-REP-01', '学生提交报修申请', '工单创建成功，状态待处理', '通过'],
        ['TC-REP-02', '管理员受理工单', '状态变为处理中', '通过'],
        ['TC-REP-03', '管理员完成工单', '状态变为已完成', '通过'],
        ['TC-REP-04', '管理员驳回工单', '状态变为已驳回', '通过'],
        ['TC-REP-05', '学生查看个人工单', '仅返回本人的工单', '通过'],
        ['TC-REP-06', '管理员查看全部工单', '返回所有工单列表', '通过'],
        ['TC-REP-07', '提交空描述报修', '返回参数校验失败', '通过'],
        ['TC-REP-08', '工单编号自动生成', '格式R+日期+6位随机数', '通过'],
    ]
)

add_paragraph('6.2.5 文件上传模块测试', bold=True, size=12)
add_table(
    ['编号', '测试用例', '预期结果', '实际结果'],
    [
        ['TC-FILE-01', '上传JPG图片', '返回图片URL', '通过'],
        ['TC-FILE-02', '上传PNG图片', '返回图片URL', '通过'],
        ['TC-FILE-03', '上传超过5MB文件', '返回文件过大错误', '通过'],
        ['TC-FILE-04', '上传非图片格式文件', '返回格式不支持错误', '通过'],
    ]
)

doc.add_heading('6.3 测试结果与分析', level=2)
add_body('经过全面的功能测试，系统各项功能均达到预期目标。认证模块的JWT Token机制和@NoAuth/@RoleRequired注解驱动的权限控制运行正常，有效保障了系统安全。报修工单状态机流转逻辑正确，各状态间转换合法。失物招领审核机制完善，管理员可正常审核和驳回信息。')
add_body('在测试过程中发现并修复了以下问题：')
add_bullet('管理后台白屏问题：修复了API模块命名导出缺失和响应数据解析错误')
add_bullet('小程序重复创建用户：修复了登录凭证策略，从临时code改为deviceId持久化')
add_bullet('Redis不可用导致500错误：配置stop-writes-on-bgsave-error=no')
add_bullet('枚举序列化问题：添加@JsonValue/@JsonCreator注解')
add_bullet('用户管理页面为空：添加数据加载和删除功能')
add_bullet('CaptchaSlider组件Bug：修复token传递和模板变量引用')
add_body('以上问题均已修复并验证通过，系统整体运行稳定，功能完整可用。')

# ============================================================
# 保存文档
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '项目开发报告-学号-姓名.docx')
doc.save(output_path)
print(f'文档已保存至: {output_path}')