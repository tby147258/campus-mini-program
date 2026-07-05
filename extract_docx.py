import docx
import os

base = r"c:\Users\24361\Desktop\实训"

for fname in ["可行性分析报告.docx", "需求分析参考示例.docx", "项目开发报告-格式模板.docx"]:
    fpath = os.path.join(base, fname)
    print("=" * 60)
    print(f"FILE: {fname}")
    print("=" * 60)
    doc = docx.Document(fpath)

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{para.style.name}] {para.text[:300]}")

    for ti, table in enumerate(doc.tables):
        print(f"\n--- Table {ti+1} ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text[:100] for cell in row.cells]
            print(f"  Row {ri}: {' | '.join(cells)}")
        if ti > 8:
            print(f"... ({len(doc.tables)} tables total, showing first 10)")
            break
    print()
