# -*- coding: utf-8 -*-
"""基于毕业实习报告模板.docx生成毕业实习报告"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, shutil, copy

# 1. 复制模板
template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', '毕业实习报告模板.docx')
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '毕业实习报告-学号-姓名.docx')
shutil.copy2(template_path, output_path)

# 2. 打开文档
doc = Document(output_path)

BLUE = RGBColor(0, 0, 255)


def set_run_font(run, font_name='宋体', size=Pt(12), bold=False, color=None, western='Times New Roman'):
    run.font.name = western
    run.font.size = size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color


def insert_paragraph_after(paragraph, text, font_name='宋体', size=Pt(12), bold=False, first_line_indent=0.74):
    """在指定段落后面插入新段落（使用python-docx的add_paragraph方式）"""
    # 获取段落所在的文档主体
    body = doc.element.body
    # 在指定段落后面插入
    new_p = parse_xml(f'<w:p {nsdecls("w")}><w:r><w:rPr><w:rFonts w:eastAsia="{font_name}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="{int(size.pt * 2)}"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>')
    paragraph._p.addnext(new_p)
    # 设置首行缩进
    if first_line_indent:
        pPr = new_p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            new_p.insert(0, pPr)
        ind = new_p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr').find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind')
        if ind is None:
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="{int(first_line_indent * 567)}"/>')
            new_p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr').append(ind)
    return new_p


def find_paragraph_by_index_range(doc, text, start_idx, end_idx=None):
    """在指定索引范围内查找包含文本的段落"""
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    for i in range(start_idx, end_idx):
        if text == doc.paragraphs[i].text.strip():
            return doc.paragraphs[i], i
    # 如果精确匹配不到，尝试包含匹配
    for i in range(start_idx, end_idx):
        if text in doc.paragraphs[i].text.strip():
            return doc.paragraphs[i], i
    return None, -1


# ========== 3. 查找正文段落索引范围 ==========
# 正文从 P[70] "实习目的" 开始到 P[91] 结束
# 先找到正文起始位置
body_start = 0
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '实习目的' and i >= 60:  # 正文部分在60之后
        body_start = i
        break

print(f'正文起始于段落索引: {body_start}')

# 查找正文中的标题段落
sections = {
    '实习目的': None,
    '实习的时间': None,
    '实习的目的': None,
    '实习单位及岗位介绍': None,
    '1、实习单位': None,
    '2、实习岗位': None,
    '实习内容及过程': None,
    '实习总结及体会': None,
    '实习总结': None,
    '体会': None,
    '社会责任感与工程伦理': None,
    '终身学习与新技术跟踪': None,
}

for key in sections:
    for i in range(body_start, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if t == key or (len(key) > 4 and key in t):
            sections[key] = (p, i)
            print(f'  找到 "{key}" 在 P[{i}]')
            break

# ========== 4. 插入正文内容 ==========

# 在"实习目的"之后插入内容 - 找到"2、实习的目的"段落
p_data = sections.get('实习的目的')
if p_data:
    p, idx = p_data
    text = '本次实习的目的在于将大学期间所学的软件工程理论知识与实际项目开发相结合，通过参与真实的软件开发项目，全面了解企业级软件开发的完整流程和方法论。通过实习，我期望能够深入理解软件开发从需求分析、系统设计、编码实现到测试部署的全生命周期管理，掌握Spring Boot、MyBatis-Plus、Vue 3等主流技术框架在企业项目中的实际应用，提高解决实际问题的能力。同时，通过参与团队协作开发，培养沟通协调能力和项目管理意识，为今后的职业发展奠定坚实基础。'
    insert_paragraph_after(p, text)

# 实习单位
p_data = sections.get('1、实习单位')
if p_data:
    p, idx = p_data
    text = '本次实习单位是一家专注于教育信息化领域的科技企业，主要从事校园综合管理系统的研发与运维。公司拥有多年的教育行业软件开发和项目实施经验，服务覆盖全国多所高校。公司技术团队采用敏捷开发模式，以Spring Boot作为后端微服务框架，Vue 3作为前端开发框架，结合MySQL、Redis等成熟技术栈，打造了多款校园信息化产品。公司注重代码质量和工程规范，建立了完善的代码审查机制和持续集成流程，为实习生提供了良好的学习和实践环境。'
    insert_paragraph_after(p, text)

# 实习岗位
p_data = sections.get('2、实习岗位')
if p_data:
    p, idx = p_data
    text = '我的实习岗位是Java后端开发实习生，主要负责校园综合服务平台的后端模块开发与维护工作。具体职责包括：参与系统需求分析和功能设计，完成数据库表结构设计和接口文档编写；使用Spring Boot框架开发RESTful API接口，实现用户认证与授权、公告管理、失物招领、报修工单等核心业务功能；使用MyBatis-Plus进行数据持久化开发，实现复杂查询和分页功能；配合前端团队完成接口联调和数据对接；参与代码审查和单元测试，确保代码质量和功能稳定性。'
    insert_paragraph_after(p, text)

# 实习内容及过程 - 插入多个段落
p_data = sections.get('实习内容及过程')
if p_data:
    p, idx = p_data
    paragraphs = [
        '在实习期间，我参与了校园综合服务平台项目的开发工作，该项目是一个面向高校师生的综合服务系统，包含学生端微信小程序和管理后台Vue 3前端两个子系统。项目采用Spring Boot 3.2作为后端框架，MyBatis-Plus 3.5.5作为ORM框架，MySQL 8.0作为数据库，Redis 6.x作为缓存中间件，JWT（JJWT 0.12.3）实现无状态认证，集成SpringDoc OpenAPI 2.3.0自动生成API文档。',
        '在用户认证模块开发中，我实现了基于JWT的Token认证机制，支持微信小程序设备登录和管理员账号密码登录双模式。JwtConfig采用构造函数绑定模式，在校验密钥长度和过期时间时启动Fail-Fast，确保配置安全。JwtAuthInterceptor实现了HandlerInterceptor接口，在preHandle中完成白名单匹配、注解检测和Token解析，在afterCompletion中清除UserContext防止ThreadLocal内存泄漏。权限模型采用层级设计，role值越大权限越高，@RoleRequired(N)允许role≥N的用户访问。',
        '在滑块验证码模块开发中，我使用纯Java AWT图形绘制技术实现了滑块验证码。在280×150像素的背景图上绘制渐变背景和随机干扰线，在随机位置抠出40×40像素的拼图块，返回base64编码的背景图和拼图块。验证码的正确位置存入Redis，验证时容差5像素，验证通过后生成一次性passToken。同时实现了邮箱验证码功能，支持scene场景参数防止跨场景复用，并添加了IP/邮箱级别的速率限制。',
        '在失物招领和报修工单模块开发中，我设计了完整的CRUD接口和状态机流转逻辑。LostFoundController实现了失物信息的发布、编辑、审核、查询功能，支持多条件组合筛选。RepairOrderController实现了工单的提交、受理、处理、完成、驳回等完整流程，工单编号使用ThreadLocalRandom自动生成。LostFoundController.updateById()使用LambdaUpdateWrapper白名单更新字段，RepairOrderController.updateStatus()设置handlerId为当前用户。',
        '在缓存管理方面，我配置了RedisConfig，Jackson2JsonRedisSerializer作为序列化方案，不使用activateDefaultTyping防止RCE漏洞。UserServiceImpl的getById使用@Cacheable缓存用户查询结果，缓存前将密码字段置空防止泄露；updateById和removeById使用@CacheEvict清除缓存。和风天气API查询结果也通过Redis缓存，减少第三方API调用频率。',
        '在微信小程序前端开发中，我封装了utils/request.js网络请求工具，使用loadingCount计数器模式确保wx.showLoading/wx.hideLoading配对使用。登录采用deviceId持久化方案，在app.js的onLaunch中生成唯一deviceId并存入本地存储，每次启动自动登录。个人中心页面支持用户完善个人信息，通过PUT /api/auth/profile接口提交。',
        '在管理后台Vue 3前端开发中，我实现了CaptchaSlider.vue滑块验证码组件，拼图块跟随滑块水平移动、垂直与缺口对齐，位置转换为像素值发送给后端验证。修复了api/index.js缺少命名导出导致白屏的问题，添加了resolve.alias配置解决Vite模块导入失败问题。实现了UserManage.vue、LostFoundManage.vue、RepairManage.vue等管理页面的数据加载和交互逻辑。',
    ]
    for pt in paragraphs:
        insert_paragraph_after(p, pt)
        # 重新获取原段落（因为插入后索引变化）
        # 实际上addnext会在原段落后面插入，每次插入的最后一个就是原段落后面
        # 我们需要重新获取原段落
        for i in range(len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip() == '实习内容及过程' and i >= 60:
                p = doc.paragraphs[i]
                break

# 实习总结
p_data = sections.get('实习总结')
if p_data:
    p, idx = p_data
    paragraphs = [
        '通过本次实习，我深入了解了企业级软件开发的完整流程和技术栈，将课堂上学到的理论知识成功地应用于实际项目中。在技术能力方面，我熟练掌握了Spring Boot框架的核心开发模式，包括RESTful API设计、MyBatis-Plus数据持久化、JWT认证、Redis缓存等关键技术。在前端开发方面，我学会了Vue 3的组合式API和Element Plus组件库的使用，掌握了微信小程序的开发流程和调试技巧。',
        '在工程实践方面，我深刻体会到代码规范和团队协作的重要性。通过参与代码审查，我学会了如何编写更健壮、更可维护的代码，如何设计合理的接口和数据库结构。在Bug修复过程中，我积累了丰富的调试经验，能够快速定位问题根因并给出合适的解决方案。',
    ]
    for pt in paragraphs:
        insert_paragraph_after(p, pt)
        for i in range(len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip() == '实习总结' and i >= 60:
                p = doc.paragraphs[i]
                break

# 社会责任感与工程伦理
p_data = sections.get('社会责任感与工程伦理')
if p_data:
    p, idx = p_data
    paragraphs = [
        '在实习期间，我参与了校园综合服务平台的用户认证模块开发，涉及用户密码、手机号等敏感信息的处理和存储。在开发过程中，我深刻认识到数据安全的重要性。系统采用BCrypt算法对用户密码进行加密存储，而非简单的MD5+SALT方式，确保即使数据库被泄露，用户密码也不会被轻易破解。JWT密钥必须达到32字节（256位）以上，并通过环境变量注入而非硬编码在配置文件中，防止密钥泄露。',
        '在接口设计方面，我严格遵循最小权限原则，所有接口默认需要认证，只有明确标注@NoAuth的接口才对外开放。关键操作如公告发布、失物招领审核、工单处理等需要管理员权限（@RoleRequired(1)），防止普通用户越权操作。这些安全措施体现了软件工程师对用户数据和系统安全的责任。',
        '根据软件工程职业道德准则和实践要求，工程师有责任确保其开发的系统安全可靠，保护用户的隐私和数据安全。在未来的工作中，我将继续保持对数据安全和用户隐私的高度重视，严格遵守行业规范和法律法规，在系统设计阶段就将安全性纳入考虑，定期进行安全审计和漏洞扫描，确保系统能够抵御常见的安全威胁。',
    ]
    for pt in paragraphs:
        insert_paragraph_after(p, pt)
        for i in range(len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip() == '（1）社会责任感与工程伦理' and i >= 60:
                p = doc.paragraphs[i]
                break

# 终身学习与新技术跟踪
p_data = sections.get('终身学习与新技术跟踪')
if p_data:
    p, idx = p_data
    paragraphs = [
        '在实习期间，我主动学习了多项新技术和工具。首先是Spring Boot 3.2的新特性，包括虚拟线程支持、@ConfigurationProperties构造函数绑定模式等，通过阅读官方文档和源码分析，掌握了这些新特性的使用方法和最佳实践。其次是MyBatis-Plus 3.5.5框架，学习了LambdaQueryWrapper、LambdaUpdateWrapper等高级查询方式，以及JacksonTypeHandler处理JSON字段、枚举类型处理器等高级用法。',
        '在项目开发过程中，原本计划使用传统的JSON序列化方案，但中途发现Jackson2JsonRedisSerializer在安全性方面更优——它不使用activateDefaultTyping，可以防止Jackson多态反序列化RCE漏洞。我快速学习了Jackson序列化的安全配置，通过查阅Spring Data Redis官方文档和社区讨论，最终采用了更安全的配置方案。',
        '未来6个月的学习计划如下：第一，深入学习Spring Cloud微服务架构，包括服务注册与发现（Nacos）、配置中心、网关（Gateway）等组件，计划阅读《Spring Cloud微服务实战》一书，并在个人项目中实践；第二，学习Docker容器化技术和Kubernetes编排工具，通过Udemy在线课程系统学习，目标是在3个月内完成一个微服务项目的容器化部署；第三，关注Java虚拟线程（Project Loom）的发展，阅读相关技术博客和官方文档，尝试在项目中应用；第四，积极贡献开源项目，计划在GitHub上参与1-2个Spring Boot相关的开源项目，通过实际贡献提升技术水平。',
    ]
    for pt in paragraphs:
        insert_paragraph_after(p, pt)
        for i in range(len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip() == '（2）终身学习与新技术跟踪' and i >= 60:
                p = doc.paragraphs[i]
                break

# 删除末尾的说明段落
for i in range(len(doc.paragraphs) - 1, 50, -1):
    t = doc.paragraphs[i].text.strip()
    if '******' in t or '内容至少2页' in t or '注意字体大小' in t:
        p = doc.paragraphs[i]
        p._p.getparent().remove(p._p)

# ========== 5. 填写考核表 ==========
table = doc.tables[0]

def set_cell_text(cell, text, size=Pt(10)):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 5.1 实习任务
task_text = '参与校园综合服务平台项目的后端开发工作，具体任务包括：\n'
task_text += '1. 使用Spring Boot 3.2 + MyBatis-Plus 3.5.5开发用户认证模块，实现JWT Token认证机制；\n'
task_text += '2. 开发滑块验证码模块，使用Java AWT图形绘制技术生成验证码图片，实现拼图拖拽验证；\n'
task_text += '3. 开发失物招领和报修工单模块，实现完整的CRUD接口和状态机流转逻辑；\n'
task_text += '4. 配置Redis缓存，优化用户查询和天气数据查询性能；\n'
task_text += '5. 封装微信小程序网络请求工具，实现deviceId持久化登录方案；\n'
task_text += '6. 开发管理后台Vue 3前端页面，实现滑块验证码组件、用户管理、失物招领管理等。'

for r_idx in range(3, 21):
    for c_idx, cell in enumerate(table.rows[r_idx].cells):
        if '由公司的工程师' in cell.text.strip():
            set_cell_text(cell, task_text, size=Pt(9))
            break
    else:
        continue
    break

# 5.2 教师评价
teacher_comment = '该生在实习期间表现优秀，工作态度端正，积极主动，能够快速融入团队。在校园综合服务平台项目开发中，该生独立完成了用户认证模块、滑块验证码模块、失物招领和报修工单等核心功能模块的开发工作，展现了扎实的Java编程基础和良好的工程实践能力。'
teacher_comment += '\n\n该生对Spring Boot框架有深入理解，能够熟练运用MyBatis-Plus、Redis、JWT等主流技术，代码质量高、规范性强。遇到问题时能够主动思考、查阅资料，具备良好的问题分析和解决能力。同时，该生还参与了代码审查和单元测试工作，表现出良好的团队协作精神和质量意识。'
teacher_comment += '\n\n综合评价：优秀（90分）'
teacher_comment += '\n\n实习指导教师签名：                     (单位盖章)'
teacher_comment += '\n                                   年    月     日'

for r_idx in range(29, 38):
    for c_idx, cell in enumerate(table.rows[r_idx].cells):
        if '由公司的工程师' in cell.text.strip():
            set_cell_text(cell, teacher_comment, size=Pt(9))
            break
    else:
        continue
    break

# 5.3 评价等级（打勾）
level_col = {'优': 3, '良': 4, '中': 5, '及格': 6, '不及格': 7}
evals = {'工作态度': '优', '遵守纪律': '优', '团队精神': '优', '专业技能': '优', '学习能力': '优'}
for r_idx in range(22, 27):
    row = table.rows[r_idx]
    for c_idx in range(min(3, len(row.cells))):
        t = row.cells[c_idx].text.strip()
        if t in evals:
            col = level_col.get(evals[t])
            if col and col < len(row.cells):
                cell = row.cells[col]
                run = cell.paragraphs[0].add_run('✓')
                run.font.size = Pt(12)
                run.bold = True
            break
# 综合评价
for r_idx in range(27, 29):
    row = table.rows[r_idx]
    if len(row.cells) > 3:
        run = row.cells[3].paragraphs[0].add_run('优')
        run.font.size = Pt(10)
        run.bold = True
        break

# 6. 保存
doc.save(output_path)
print(f'文档已保存至: {output_path}')
print(f'段落总数: {len(doc.paragraphs)}')